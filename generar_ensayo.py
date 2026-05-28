import os
import sys
import urllib.request
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_number(run):
    """Inserta el campo XML para el número de página en Word."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def descargar_logo():
    """Descarga el logo de la UNAH de Wikimedia Commons. Si falla, retorna None."""
    url = "https://upload.wikimedia.org/wikipedia/commons/thumb/a/a6/Escudo_de_la_UNAH.svg/960px-Escudo_de_la_UNAH.svg.png"
    dest = "logo_unah.png"
    print("Intentando descargar el logotipo institucional de la UNAH...")
    try:
        req = urllib.request.Request(
            url, 
            headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'}
        )
        with urllib.request.urlopen(req, timeout=10) as response, open(dest, 'wb') as out_file:
            out_file.write(response.read())
        print("Logotipo descargado con éxito.")
        return dest
    except Exception as e:
        print(f"No se pudo descargar el logotipo (Error: {e}). Se generará el documento sin la imagen (dejando espacio).")
        return None

def add_body_paragraph(doc, text):
    """Agrega un párrafo de cuerpo del texto con formato APA 7 estricto (doble espacio, sangría, color negro)."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_heading_1(doc, text):
    """Agrega un encabezado de Nivel 1 (Centrado, negrita, color negro)."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_heading_2(doc, text):
    """Agrega un encabezado de Nivel 2 (Alineado a la izquierda, negrita, color negro)."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_reference_paragraph(doc, runs_data):
    """Agrega una referencia bibliográfica con sangría francesa (1.27 cm) y doble espacio.
       runs_data es una lista de tuplas (texto, es_cursiva)"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for text, is_italic in runs_data:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.italic = is_italic
    return p

def main():
    doc = Document()
    
    # 1. Configuración general de márgenes (2.54 cm en todos los lados)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        
        # Habilitar numeración de página en el encabezado
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_header = header_para.add_run()
        run_header.font.name = 'Times New Roman'
        run_header.font.size = Pt(12)
        run_header.font.color.rgb = RGBColor(0, 0, 0)
        add_page_number(run_header)

    # ------------------ PORTADA (ESTILO UNAH CENTRADO) ------------------
    # Títulos superiores de la universidad
    p_unah = doc.add_paragraph()
    p_unah.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_unah.paragraph_format.line_spacing = 1.15
    p_unah.paragraph_format.space_before = Pt(0)
    p_unah.paragraph_format.space_after = Pt(6)
    r_unah = p_unah.add_run("UNIVERSIDAD NACIONAL AUTÓNOMA DE HONDURAS\n")
    r_unah.font.name = 'Times New Roman'
    r_unah.font.size = Pt(12)
    r_unah.bold = True
    r_unah.font.color.rgb = RGBColor(0, 0, 0)
    
    r_fac = p_unah.add_run("FACULTAD DE CIENCIAS ECONÓMICAS, ADMINISTRATIVAS Y CONTABLES\n")
    r_fac.font.name = 'Times New Roman'
    r_fac.font.size = Pt(11)
    r_fac.bold = True
    r_fac.font.color.rgb = RGBColor(0, 0, 0)
    
    r_lic = p_unah.add_run("LICENCIATURA EN INFORMÁTICA ADMINISTRATIVA")
    r_lic.font.name = 'Times New Roman'
    r_lic.font.size = Pt(12)
    r_lic.bold = True
    r_lic.font.color.rgb = RGBColor(0, 0, 0)

    # Espacio para el logotipo
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(12)
    p_logo.paragraph_format.space_after = Pt(12)
    
    logo_path = descargar_logo()
    if logo_path and os.path.exists(logo_path):
        try:
            p_logo.add_run().add_picture(logo_path, width=Inches(1.8))
        except Exception as e:
            print(f"Error al insertar la imagen: {e}. Colocando marcador de posición.")
            p_logo.add_run("[Escudo de la UNAH]").font.color.rgb = RGBColor(0,0,0)
    else:
        p_logo.add_run("\n\n[Escudo de la UNAH]\n\n").font.color.rgb = RGBColor(0,0,0)

    # Detalles de la Clase
    p_clase = doc.add_paragraph()
    p_clase.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_clase.paragraph_format.line_spacing = 1.3
    p_clase.paragraph_format.space_before = Pt(12)
    p_clase.paragraph_format.space_after = Pt(18)
    
    r_clase = p_clase.add_run("CLASE: DIA-127- ARQUITECTURA DE SOFTWARE I.\nII - PAC 2026\nTELEDOCENCIA\n\nACTIVIDAD #: 01 Análisis de Caso Real de Disrupción\n")
    r_clase.font.name = 'Times New Roman'
    r_clase.font.size = Pt(12)
    r_clase.bold = True
    r_clase.font.color.rgb = RGBColor(0, 0, 0)

    # Tabla con datos del estudiante (Alineado y ordenado como el original)
    p_sust = doc.add_paragraph()
    p_sust.paragraph_format.space_before = Pt(12)
    p_sust.paragraph_format.space_after = Pt(6)
    r_sust = p_sust.add_run("SUSTENTADA POR:")
    r_sust.font.name = 'Times New Roman'
    r_sust.font.size = Pt(12)
    r_sust.bold = True
    r_sust.font.color.rgb = RGBColor(0, 0, 0)
    
    # Creamos una tabla sin bordes para estructurar los datos perfectamente
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Estilo sin bordes
    table.style = 'Normal Table'
    
    headers = ["NOMBRE:", "NÚMERO DE CUENTA:", "GRUPO #:"]
    data = ["PAOLO KEVIN BRANT SIERRA", "20042000562", "1"]
    
    # Rellenar encabezados de la tabla
    for i, head_text in enumerate(headers):
        cell = table.cell(0, i)
        cell.width = Inches(2.2)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(head_text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
        
    # Rellenar datos
    for i, data_text in enumerate(data):
        cell = table.cell(1, i)
        cell.width = Inches(2.2)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(data_text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        r.bold = True
        r.italic = True
        r.font.color.rgb = RGBColor(0, 0, 0)

    # Capítulo, Unidad y Catedrático
    p_cat = doc.add_paragraph()
    p_cat.paragraph_format.line_spacing = 1.3
    p_cat.paragraph_format.space_before = Pt(24)
    p_cat.paragraph_format.space_after = Pt(24)
    
    # Tabla simple de 2 columnas para alinear CAPITULO, UNIDAD y CATEDRÁTICO
    table_cat = doc.add_table(rows=3, cols=2)
    table_cat.style = 'Normal Table'
    
    labels = ["CAPITULO:", "UNIDAD:", "CATEDRÁTICO:"]
    values = ["1.", "I.", "PAOLO BRANT."]
    
    for i in range(3):
        # Etiqueta
        cell_lbl = table_cat.cell(i, 0)
        cell_lbl.width = Inches(1.8)
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(labels[i])
        r_lbl.font.name = 'Times New Roman'
        r_lbl.font.size = Pt(12)
        r_lbl.bold = True
        r_lbl.font.color.rgb = RGBColor(0, 0, 0)
        
        # Valor
        cell_val = table_cat.cell(i, 1)
        cell_val.width = Inches(4.5)
        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(values[i])
        r_val.font.name = 'Times New Roman'
        r_val.font.size = Pt(12)
        r_val.bold = True
        r_val.font.color.rgb = RGBColor(0, 0, 0)

    # Fecha al pie de página (San Pedro Sula)
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(36)
    p_date.paragraph_format.space_after = Pt(0)
    r_date = p_date.add_run("SAN PEDRO SULA, HONDURAS, 11 DE MAYO DE 2026.")
    r_date.font.name = 'Times New Roman'
    r_date.font.size = Pt(12)
    r_date.bold = True
    r_date.font.color.rgb = RGBColor(0, 0, 0)

    # Salto de página después de la portada
    doc.add_page_break()

    # ------------------ CUERPO DEL ENSAYO (PÁGINAS 2 Y 3) ------------------
    # Título del Ensayo
    add_heading_1(doc, "El Colapso de Sun Microsystems bajo la Óptica de la Curva en S: Un Análisis Estratégico a través de los Marcos de Day y Laudon")
    
    # Párrafo 1: Introducción
    add_body_paragraph(
        doc,
        "En la historia de la industria tecnológica, el auge y caída de Sun Microsystems representa uno de los casos de estudio más "
        "ilustrativos sobre la dinámica de la innovación y la inercia estratégica. Fundada en 1982, la compañía se convirtió rápidamente "
        "en un gigante tecnológico al dominar el mercado de las estaciones de trabajo y los servidores corporativos de alto rendimiento "
        "mediante una arquitectura integrada verticalmente que combinaba su procesador SPARC y el sistema operativo Solaris (basado en Unix). "
        "Durante la época del auge del Internet y la burbuja de las empresas puntocom en los años noventa, Sun Microsystems experimentó un "
        "crecimiento sin precedentes bajo el lema institucional de que la red era la computadora. Sin embargo, su posición de liderazgo absoluto "
        "ocultó una complacencia que le impidió anticipar y responder al surgimiento de nuevas tecnologías disruptivas. Este ensayo analiza "
        "el colapso estratégico de Sun Microsystems desde la teoría de la Curva en S de la tecnología y aplica los marcos teóricos de Day (2007) "
        "y Laudon y Laudon (2020) para explicar cómo la inercia organizacional y las fallas en la evaluación del riesgo impidieron a la dirección de "
        "la compañía migrar oportunamente hacia la nueva curva tecnológica de los servidores de bajo costo basados en la arquitectura x86 "
        "y el sistema operativo de código abierto Linux."
    )

    # Encabezado 2: La Curva en S y la Transición Tecnológica
    add_heading_2(doc, "La Curva en S y la Transición Tecnológica")
    
    # Párrafo 2
    add_body_paragraph(
        doc,
        "La teoría de la Curva en S describe la evolución del rendimiento de una tecnología a lo largo del tiempo en relación con el esfuerzo "
        "o la inversión realizada. En sus etapas iniciales, el avance es lento debido a la falta de madurez de la tecnología; posteriormente, "
        "entra en una fase de crecimiento exponencial a medida que se refina y se adoptan mejores prácticas de ingeniería; finalmente, alcanza "
        "un límite físico y entra en una meseta de madurez. En el caso de Sun Microsystems, su tecnología propietaria basada en microprocesadores "
        "SPARC y el sistema operativo Solaris representaba una curva en S madura y altamente optimizada. Ofrecía un rendimiento excepcional "
        "para tareas de cómputo intensivo y bases de datos corporativas gigantescas, justificando sus elevados precios y garantizando a la empresa "
        "márgenes brutos superiores al cincuenta por ciento."
    )
    
    # Párrafo 3
    add_body_paragraph(
        doc,
        "Por otro lado, la curva tecnológica emergente a finales de los noventa estaba definida por los procesadores de arquitectura abierta "
        "x86 (fabricados por Intel y AMD) combinados con sistemas operativos basados en Linux. Al principio, esta tecnología se encontraba en la "
        "base de su propia curva en S: su rendimiento era muy inferior, carecía del soporte de nivel empresarial necesario para aplicaciones "
        "críticas y los sistemas de base de datos grandes no se ejecutaban de manera eficiente sobre ella. Sin embargo, la trayectoria de la curva "
        "x86/Linux avanzaba a una velocidad sustancialmente mayor debido a la estandarización del hardware y la colaboración masiva del ecosistema "
        "de código abierto. A medida que esta nueva tecnología ascendía en su curva de rendimiento, comenzó a cruzar el umbral de lo que era aceptable "
        "para las empresas medianas y grandes, ofreciendo una alternativa con una relación costo-beneficio infinitamente superior a la de los "
        "costosos servidores propietarios de Sun Microsystems."
    )

    # Encabezado 3: El Marco de Day y la Evaluación de Riesgos
    add_heading_2(doc, "El Marco de Day y la Evaluación de Riesgos")
    
    # Párrafo 4
    add_body_paragraph(
        doc,
        "El marco teórico de George Day sobre la gestión de la innovación y la toma de decisiones directivas nos permite comprender por qué los "
        "ejecutivos de Sun Microsystems, liderados por Scott McNealy, se negaron a saltar a la nueva curva tecnológica. Según Day (2007), las organizaciones "
        "a menudo sufren de una incapacidad estratégica para evaluar adecuadamente el riesgo de omisión (no adoptar una nueva tecnología) en comparación "
        "con el riesgo de comisión (el riesgo financiero asociado con realizar una inversión fallida). La directiva de Sun Microsystems estaba atrapada "
        "en un sesgo cognitivo provocado por su propio éxito. La introducción de servidores de bajo costo basados en x86 para competir en el mercado "
        "de gama baja representaba un dilema de canibalización: comercializar hardware de bajo margen amenazaba directamente las ventas de su línea "
        "premium de servidores SPARC, que financiaba su enorme presupuesto de investigación y desarrollo."
    )
    
    # Párrafo 5
    add_body_paragraph(
        doc,
        "Los directivos evaluaron la transición tecnológica bajo una perspectiva puramente defensiva, optando por ignorar la amenaza x86/Linux "
        "con la convicción de que los clientes corporativos jamás migrarían sus sistemas centrales a arquitecturas no propietarias e inseguras. "
        "Esta subestimación del ritmo de mejora de la curva emergente y la sobrevaloración de la lealtad del cliente crearon un punto ciego estratégico. "
        "De acuerdo con el marco de Day (2007), la incapacidad de la empresa para redefinir su mercado y su resistencia a sacrificar los márgenes a corto "
        "plazo para asegurar su viabilidad a largo plazo aceleraron su declive institucional."
    )

    # Encabezado 4: El Marco de Laudon: Inercia Organizacional y Sistemas de Información
    add_heading_2(doc, "El Marco de Laudon: Inercia Organizacional y Sistemas de Información")
    
    # Párrafo 6
    add_body_paragraph(
        doc,
        "Desde la perspectiva de Laudon y Laudon (2020), los Sistemas de Información Estratégicos (SIE) deben alinearse con la estructura "
        "organizacional, la cultura y los procesos de negocio para sostener la ventaja competitiva. El colapso de Sun Microsystems ilustra "
        "de manera contundente cómo la inercia organizacional y la falta de alineación entre la arquitectura de software estratégica y los "
        "sistemas internos bloquearon la capacidad de respuesta de la compañía. El modelo de Laudon y Laudon (2020) sostiene que las organizaciones "
        "son sistemas complejos con una inercia institucional arraigada, destacando que \"las rutinas organizacionales, la política y la cultura dificultan la adopción de nuevas tecnologías\" (p. 95). La cultura de Sun Microsystems estaba fundamentada en el orgullo de su "
        "superioridad en ingeniería de hardware y sistemas propietarios. Sus sistemas de información de ventas, incentivos y distribución interna "
        "estaban diseñados exclusivamente para promover la venta de sistemas cerrados con grandes márgenes de ganancia."
    )
    
    # Párrafo 7
    add_body_paragraph(
        doc,
        "Cuando la señal del mercado comenzó a cambiar de forma dramática tras el estallido de la burbuja puntocom en el año 2000 —cuando las empresas "
        "exigieron eficiencia de costos y escalabilidad horizontal (scale-out) con servidores x86 en lugar de la escalabilidad vertical (scale-up) de "
        "Sun—, los canales internos de retroalimentación de la empresa no pudieron digerir la información. La inercia de la estructura organizacional "
        "impidió la reconfiguración de los sistemas de información corporativos para respaldar una estrategia de volumen y servicios de software. "
        "Aunque Sun Microsystems desarrolló Java e intentó liberar tarde el código de Solaris (creando OpenSolaris en 2005), la rigidez estructural "
        "de su modelo de negocio impidió la monetización efectiva del software y el soporte, evidenciando que las organizaciones no pueden cambiar "
        "su infraestructura tecnológica estratégica de forma aislada sin transformar sus procesos organizacionales subyacentes."
    )

    # Encabezado 5: Contraste de Márgenes y la Incertidumbre del Mercado Emergente
    add_heading_2(doc, "Contraste de Márgenes y la Incertidumbre del Mercado Emergente")
    
    # Párrafo 8
    add_body_paragraph(
        doc,
        "La reticencia a saltar de una curva en S a otra se explica de forma fundamental por el contraste económico entre ambos modelos. La tecnología "
        "madura de Sun Microsystems ofrecía una certidumbre operativa absoluta. Los clientes corporativos estaban dispuestos a pagar cientos de miles "
        "de dólares por un solo mainframe Unix debido a su confiabilidad y al servicio postventa personalizado. Los márgenes de ganancia establecidos "
        "permitían a Sun sostener una estructura corporativa costosa y un modelo de integración vertical. En contraste, la tecnología emergente de "
        "servidores basados en Linux y microprocesadores x86 estándar presentaba una enorme incertidumbre a principios de la década de 2000. Los márgenes "
        "de ganancia unitarios en el hardware de servidores x86 eran extremadamente reducidos, y el modelo de negocio requería un volumen de ventas "
        "masivo para ser rentable."
    )
    
    # Párrafo 9
    add_body_paragraph(
        doc,
        "Adicionalmente, el ecosistema de Linux carecía de una entidad corporativa única que garantizara la estabilidad a largo plazo del sistema, lo que "
        "generaba desconfianza en los departamentos de tecnologías de la información tradicionales. Enfrentados a esta disyuntiva, los directivos de Sun "
        "Microsystems tomaron una decisión racional a corto plazo pero catastrófica a largo plazo: decidieron proteger sus vacas lecheras de alto margen "
        "y posponer el desarrollo serio de servidores genéricos. Al hacerlo, permitieron que competidores como Hewlett-Packard, Dell e IBM capturaran "
        "y dominaran el mercado emergente. Cuando la tecnología x86 y Linux maduró lo suficiente como para competir en el segmento corporativo de "
        "misión crítica, Sun Microsystems ya no tenía el flujo de caja necesario para financiar la transición y recuperar la ventaja perdida."
    )

    # Encabezado 6: Conclusiones
    add_heading_2(doc, "Conclusiones")
    
    # Párrafo 10
    add_body_paragraph(
        doc,
        "El destino final de Sun Microsystems, adquirida por Oracle Corporation en el año 2010 por una fracción de lo que llegó a valer en su cúspide, "
        "ofrece una lección fundamental en la gestión de la arquitectura de software e infraestructura tecnológica. El colapso de este gigante demuestra "
        "que el dominio técnico en una curva en S madura no garantiza la supervivencia ante cambios de paradigma. La aplicación de los marcos de Day "
        "y Laudon evidencia que los fracasos estratégicos rara vez son causados por limitaciones de ingeniería, sino por rigideces organizacionales "
        "y sesgos directivos en la evaluación del riesgo y los márgenes financieros. Para los arquitectos de software contemporáneos y los gestores de "
        "tecnología, la lección es clara: es indispensable diseñar arquitecturas de software flexibles que permitan transicionar entre plataformas "
        "y vigilar constantemente las curvas en S emergentes, reconociendo que la inercia del modelo de negocio maduro puede ser el mayor obstáculo "
        "para la innovación disruptiva."
    )

    # Salto de página para las referencias
    doc.add_page_break()

    # ------------------ REFERENCIAS (PÁGINA DE REFERENCIAS APA 7) ------------------
    add_heading_1(doc, "Referencias")
    
    # Referencias individuales con sangría francesa
    add_reference_paragraph(
        doc,
        [
            ("Day, G. S. (2007). Is it real? Can we win? Is it worth doing? Managing risk and reward in an innovation portfolio. ", False),
            ("Harvard Business Review", True),
            (", 85(12), 110-120.", False)
        ]
    )
    add_reference_paragraph(
        doc,
        [
            ("Laudon, K. C., & Laudon, J. P. (2020). ", False),
            ("Sistemas de información gerencial", True),
            (" (16a ed.). Pearson Educación.", False)
        ]
    )
    add_reference_paragraph(
        doc,
        [
            ("Vance, A. (2009, 18 de marzo). Sun Microsystems: A technology giant's rise and fall. ", False),
            ("The New York Times", True),
            (". https://www.nytimes.com", False)
        ]
    )

    # Guardar documento
    output_filename = "Ensayo_Sun_Microsystems_Curva_S.docx"
    doc.save(output_filename)
    print(f"Documento de Word generado exitosamente: {output_filename}")

if __name__ == "__main__":
    main()
